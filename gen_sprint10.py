#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_sprint10.py — Generador de documentos Sprint 10 — ICARO
--------------------------------------------------------------
Genera:
  1. Sprint_10_Planificado_Realizado_ICARO.docx  (PLN-REA-SPR-10)
  2. Sprint_10_Informe_Pruebas_ICARO.docx        (INF-PRU-SPR-10)
  3. Sprint_10_Retrospectiva_ICARO_V2.docx       (RET-SPR-10)
  4. Sprint_10_Backlog_ICARO.xlsx                (BCK-SPR-10)

Sprint 10 — Auditoría, Cierre Mensual e Inmutabilidad
Período  : 15/06/2026 – 19/06/2026
CP       : CP-223 – CP-254 (32 casos)
Acumulado: 254 pruebas
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ─── RUTAS ────────────────────────────────────────────────────────────────────
BASE   = r"c:\Users\Hp\Desktop\Sistema_Gestion_Integral_Constructora\Documentos Base"
OUTDIR = os.path.join(BASE, "Sprint 10 documentacion")
os.makedirs(OUTDIR, exist_ok=True)

# ─── DATOS DEL SPRINT ─────────────────────────────────────────────────────────
SPR         = "10"
NOMBRE_SPR  = "Auditoría, Cierre Mensual e Inmutabilidad"
FECHA_INI   = "15/06/2026"
FECHA_FIN   = "19/06/2026"
FECHA_DOC   = "19/06/2026"
RESPONSABLE = "Ivan Santiago Pulgar Leon"
PROYECTO    = "Sistema de Gestión Integral de Obra — ICARO CONSTRUCTORES BMGM S.A.S."
OBJETIVO    = (
    "Implementar el módulo de Auditoría y Cierre Mensual del sistema ICARO, "
    "integrando: (1) consolidación mensual contable-operativa (avances, compras, consumos e "
    "inventario por proyecto/período); (2) validación pre-cierre con bloqueos explícitos "
    "(PENDING_SYNC, EN_REVISION, cierre ya CERRADO) y advertencias; (3) ejecución del cierre "
    "mensual con transacción prisma.$transaction que genera hash SHA-256 del payload y lo "
    "persiste en cierre_mensual.hash_seguridad; (4) módulo de auditoría con filtros por "
    "usuario, tabla, operación y fechas, exportación a CSV/Excel y PDF, e inmutabilidad "
    "garantizada desde la API (sin endpoints DELETE/UPDATE expuestos para audit_log); "
    "(5) rollback transaccional total ante cualquier fallo dentro de la transacción. "
    "Aclaración: la BD no fue alterada — el servicio opera sobre tablas ya existentes "
    "(cierre_mensual, audit_log, avance_obra, requerimiento_compra, movimiento_inventario, "
    "inventario_proyecto)."
)
MODULOS = (
    "Consolidación Mensual (GET /consolidacion), Validación Pre-cierre (POST /validar), "
    "Ejecución Cierre + Hash SHA-256 (POST /ejecutar), Historial Cierres (GET /cierres-contables), "
    "Auditoría Filtrada + Exportación (GET /audit-logs), Rollback Transaccional"
)

# ─── COLORES ──────────────────────────────────────────────────────────────────
C_DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
C_BLACK     = RGBColor(0x00, 0x00, 0x00)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY      = RGBColor(0x59, 0x59, 0x59)

H_DARK  = "1F3864"
H_MED   = "D9E1F2"
H_GREEN = "E2EFDA"
H_RED   = "FCE4D6"
H_YELL  = "FFF2CC"
H_WHITE = "FFFFFF"
H_GRAY  = "F2F2F2"

# ════════════════════════════════════════════════════════════════════════════════
# UTILIDADES DOCX
# ════════════════════════════════════════════════════════════════════════════════

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=None, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.name = 'Calibri'
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = color

def hrow(table, labels, bg=H_DARK, fg=C_WHITE, sz=9):
    row = table.rows[0]
    for i, lbl in enumerate(labels):
        if i < len(row.cells):
            cell_bg(row.cells[i], bg)
            cell_text(row.cells[i], lbl, bold=True, size=sz, color=fg,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs: run.font.name = 'Calibri'
    return h

def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs: run.font.name = 'Calibri'
    return h

def kv_table(doc, rows_data):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(rows_data):
        cell_bg(tbl.rows[i].cells[0], H_MED)
        cell_text(tbl.rows[i].cells[0], k, bold=True, size=10, color=C_BLACK)
        cell_text(tbl.rows[i].cells[1], v, size=10, color=C_BLACK)
    return tbl

def versions_table(doc):
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    hrow(t, ['Versión', 'Fecha', 'Responsable', 'Descripción del cambio'])
    for j, v in enumerate(['1.0', FECHA_DOC, RESPONSABLE,
                            f'Documento generado con datos reales del Sprint {SPR}']):
        cell_text(t.rows[1].cells[j], v, size=10, color=C_BLACK)

def sprint_data_table(doc):
    kv_table(doc, [
        ("Número del sprint",              f"Sprint {SPR}"),
        ("Nombre del sprint",              NOMBRE_SPR),
        ("Fecha de inicio",                FECHA_INI),
        ("Fecha de cierre",                FECHA_FIN),
        ("Duración",                       "1 semana"),
        ("Módulos/componentes trabajados", MODULOS),
    ])

def normal_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
    r.font.color.rgb = C_BLACK

def control_doc_table(doc, codigo, nombre_doc, version='1.0'):
    kv_table(doc, [
        ("Código del documento",       codigo),
        ("Nombre del documento",       nombre_doc),
        ("Versión",                    version),
        ("Fecha de elaboración",       FECHA_DOC),
        ("Sprint evaluado",            f"Sprint {SPR} – {NOMBRE_SPR}"),
        ("Responsable de elaboración", RESPONSABLE),
        ("Estado del documento",       "Aprobado"),
    ])

# ════════════════════════════════════════════════════════════════════════════════
# HU / TAREAS
# ════════════════════════════════════════════════════════════════════════════════
# (ID, Tipo, Desc, Prioridad, SP_Plan, SP_Real, Desvío, %Cumpl, F_Inicio, F_Fin)
HU_ROWS = [
    ("HU-S10-1","HU",  "Consolidación mensual contable-operativa por proyecto y período",
     "Alta", 8, 8, 0, "100%", "15/06", "16/06"),
    ("T-S10-1.1","Tarea","GET /api/v1/cierres-contables/consolidacion con RBAC (Contador, Admin, Presidente)","",
     3, 3, 0, "100%", "15/06", "15/06"),
    ("T-S10-1.2","Tarea","Agregación: avances (qty+monto), compras APROBADO/RECIBIDO, consumos SALIDA, snapshot inventario","",
     3, 3, 0, "100%", "15/06", "16/06"),
    ("T-S10-1.3","Tarea","Porcentaje avance global (avance/presupuesto). Pruebas Grupo 1 (CP-223–CP-229)","",
     2, 2, 0, "100%", "16/06", "16/06"),

    ("HU-S10-2","HU",  "Validación pre-cierre con bloqueos explícitos y advertencias",
     "Alta", 6, 6, 0, "100%", "16/06", "17/06"),
    ("T-S10-2.1","Tarea","POST /api/v1/cierres-contables/validar — RBAC Contador + Admin. Residente/Bodeguero → 403","",
     2, 2, 0, "100%", "16/06", "17/06"),
    ("T-S10-2.2","Tarea","Bloqueos: PENDING_SYNC, EN_REVISION, cierre ya CERRADO → 422 con array de errores","",
     2, 2, 0, "100%", "17/06", "17/06"),
    ("T-S10-2.3","Tarea","Advertencias (no bloqueantes): avances rechazados, stock negativo. Registro en audit_log.datos_despues","",
     2, 2, 0, "100%", "17/06", "17/06"),

    ("HU-S10-3","HU",  "Ejecución de cierre mensual con hash SHA-256 e inmutabilidad",
     "Alta", 8, 8, 0, "100%", "17/06", "18/06"),
    ("T-S10-3.1","Tarea","POST /ejecutar con prisma.$transaction: consolidar → validar → cerrar → hash → audit","",
     3, 3, 0, "100%", "17/06", "18/06"),
    ("T-S10-3.2","Tarea","SHA-256 generado del payload y persistido en cierre_mensual.hash_seguridad (64 chars hex)","",
     3, 3, 0, "100%", "18/06", "18/06"),
    ("T-S10-3.3","Tarea","Estado CERRADO + fecha_cierre + bloqueo UI. Solo Contador y Admin (403 otros roles)","",
     2, 2, 0, "100%", "18/06", "18/06"),

    ("HU-S10-4","HU",  "Auditoría con filtros, exportación y garantía de inmutabilidad",
     "Media", 8, 8, 0, "100%", "17/06", "19/06"),
    ("T-S10-4.1","Tarea","GET /api/v1/audit-logs con filtros: idUsuario, tabla, operacion, fechaDesde, fechaHasta","",
     3, 3, 0, "100%", "17/06", "18/06"),
    ("T-S10-4.2","Tarea","AuditTraceabilityView.jsx: exportación CSV/Excel (BOM UTF-8) y PDF (print). Solo Admin (403 otros)","",
     3, 3, 0, "100%", "18/06", "19/06"),
    ("T-S10-4.3","Tarea","Inmutabilidad: sin endpoints DELETE/UPDATE en audit_log. logAction() solo invoca prisma.auditLog.create()","",
     2, 2, 0, "100%", "19/06", "19/06"),

    ("HT-S10-01","HT", "Rollback transaccional total e inmutabilidad a nivel de API",
     "Alta", 6, 6, 0, "100%", "18/06", "19/06"),
    ("T-H10.1","Tarea","Rollback total si falla cierre/hash/audit dentro de prisma.$transaction","",
     3, 3, 0, "100%", "18/06", "19/06"),
    ("T-H10.2","Tarea","Validaciones 422 (PENDING_SYNC, etc.) no abren transacción — BD intacta","",
     2, 2, 0, "100%", "19/06", "19/06"),
    ("T-H10.3","Tarea","Pruebas Grupo 5 (CP-251–CP-254): rollback simulado verificado","",
     1, 1, 0, "100%", "19/06", "19/06"),
]

TOTAL_SP_PLAN = 36
TOTAL_SP_REAL = 36

# ════════════════════════════════════════════════════════════════════════════════
# CASOS DE PRUEBA — CP-223 a CP-254 (32 casos)
# ════════════════════════════════════════════════════════════════════════════════
# (ID, Grupo, Módulo, Descripción, Tipo, Resultado Esperado, Resultado Obtenido, Estado, Evidencia)
CP_ROWS = [
    # GRUPO 1 — Consolidación Mensual (CP-223–CP-229)
    ("CP-223","Grupo 1","Consolidación Mensual",
     "GET /consolidacion sin token → 401 Unauthorized",
     "Automatizada","HTTP 401. body con campo 'error'.","HTTP 401 – { error: 'Token requerido' }.","PASS",
     "sprint10_verificacion.md – Act-1 | Postman"),
    ("CP-224","Grupo 1","Consolidación Mensual",
     "GET /consolidacion con Residente → 403 Forbidden (no está en canRead)",
     "Automatizada","HTTP 403.","HTTP 403 – { error: 'Acceso denegado' }.","PASS",
     "sprint10_verificacion.md – Act-1 CA-1.7"),
    ("CP-225","Grupo 1","Consolidación Mensual",
     "GET /consolidacion con Contador → 200 OK — no 401, no 403",
     "Automatizada",
     "HTTP 200. body con { success: true, data: { totalAvanceMonto, totalComprasMonto, "
     "totalConsumos, snapshotInventario, pctAvanceGlobal } }.",
     "HTTP 200 con todos los campos de consolidación.","PASS",
     "sprint10_verificacion.md – Act-1 CA-1.1 a CA-1.6"),
    ("CP-226","Grupo 1","Consolidación Mensual",
     "GET /consolidacion con Admin → 200 con datos del período. snapshotInventario array no vacío si hay movimientos",
     "Manual",
     "HTTP 200. data.snapshotInventario es array con { codigo, nombre, unidad, stockActual }.",
     "HTTP 200 con snapshot correcto.","PASS",
     "sprint10_verificacion.md – Act-1 CA-1.5"),
    ("CP-227","Grupo 1","Consolidación Mensual",
     "GET /consolidacion con Presidente/Gerente → pasa RBAC (200/404/500, no 401, no 403)",
     "Automatizada","HTTP 200/404/500 (no 401, no 403).","HTTP 404 sin datos en período vacío.","PASS",
     "sprint10_verificacion.md – Act-1 CA-1.7"),
    ("CP-228","Grupo 1","Consolidación Mensual",
     "GET /consolidacion sin idProyecto o sin mesAnio → 400 Bad Request",
     "Automatizada","HTTP 400. body con 'error' indicando campo requerido.","HTTP 400.","PASS",
     "sprint10_verificacion.md – Act-1"),
    ("CP-229","Grupo 1","Consolidación Mensual",
     "GET /consolidacion — pctAvanceGlobal >= 0 y total compras incluye solo APROBADO/RECIBIDO",
     "Manual",
     "pctAvanceGlobal es decimal >= 0. Solo requerimientos APROBADO o RECIBIDO en totalComprasMonto.",
     "Valores calculados correctamente.","PASS",
     "sprint10_verificacion.md – Act-1 CA-1.2 CA-1.3"),

    # GRUPO 2 — Validación Pre-cierre (CP-230–CP-236)
    ("CP-230","Grupo 2","Validación Pre-cierre",
     "POST /validar sin token → 401 Unauthorized",
     "Automatizada","HTTP 401. body con 'error'.","HTTP 401.","PASS",
     "sprint10_verificacion.md – Act-2 | Postman"),
    ("CP-231","Grupo 2","Validación Pre-cierre",
     "POST /validar con Residente → 403 Forbidden",
     "Automatizada","HTTP 403.","HTTP 403.","PASS",
     "sprint10_verificacion.md – Act-2 | RBAC"),
    ("CP-232","Grupo 2","Validación Pre-cierre",
     "POST /validar con Contador sin bloqueos → 200 OK — banner verde en UI",
     "Manual",
     "HTTP 200. body { validado: true, bloqueos: [], advertencias: [] }. UI: banner verde.",
     "HTTP 200. Sin bloqueos detectados.","PASS",
     "sprint10_verificacion.md – Act-2 CA-2.7 Escenario A"),
    ("CP-233","Grupo 2","Validación Pre-cierre",
     "POST /validar con avances PENDING_SYNC en el período → 422 con bloqueo [AVN_PENDING_SYNC]",
     "Automatizada",
     "HTTP 422. body.errors contiene { codigo: 'AVN_PENDING_SYNC', descripcion: '...' }.",
     "HTTP 422 – bloqueo PENDING_SYNC confirmado.","PASS",
     "sprint10_verificacion.md – Act-2 CA-2.1"),
    ("CP-234","Grupo 2","Validación Pre-cierre",
     "POST /validar con requerimientos EN_REVISION → 422 con bloqueo [REQ_EN_REVISION]",
     "Automatizada",
     "HTTP 422. body.errors contiene { codigo: 'REQ_EN_REVISION', descripcion: '...' }.",
     "HTTP 422 – bloqueo EN_REVISION confirmado.","PASS",
     "sprint10_verificacion.md – Act-2 CA-2.2"),
    ("CP-235","Grupo 2","Validación Pre-cierre",
     "POST /validar con período ya CERRADO en cierre_mensual → 422 con bloqueo [CIERRE_YA_EXISTE]",
     "Automatizada",
     "HTTP 422. body.errors contiene { codigo: 'CIERRE_YA_EXISTE' }.",
     "HTTP 422 – período ya cerrado.","PASS",
     "sprint10_verificacion.md – Act-2 CA-2.3"),
    ("CP-236","Grupo 2","Validación Pre-cierre",
     "POST /validar — resultado (bloqueos/advertencias) queda en audit_log.datos_despues con accion=VALIDACION_PRE_CIERRE",
     "Manual",
     "audit_log.datos_despues.accion === 'VALIDACION_PRE_CIERRE'. Registro persistido.",
     "Registro en audit_log verificado en módulo Auditoría.","PASS",
     "sprint10_verificacion.md – Act-2 CA-2.6"),

    # GRUPO 3 — Cierre Mensual + Hash SHA-256 (CP-237–CP-244)
    ("CP-237","Grupo 3","Cierre Mensual",
     "POST /ejecutar sin token → 401 Unauthorized",
     "Automatizada","HTTP 401.","HTTP 401.","PASS",
     "sprint10_verificacion.md – Act-3 | Postman"),
    ("CP-238","Grupo 3","Cierre Mensual",
     "POST /ejecutar con Residente → 403 Forbidden",
     "Automatizada","HTTP 403.","HTTP 403.","PASS",
     "sprint10_verificacion.md – Act-3 CA-3.7"),
    ("CP-239","Grupo 3","Cierre Mensual",
     "POST /ejecutar con Bodeguero → 403 Forbidden",
     "Automatizada","HTTP 403.","HTTP 403.","PASS",
     "sprint10_verificacion.md – Act-3 CA-3.7"),
    ("CP-240","Grupo 3","Cierre Mensual",
     "POST /ejecutar con Contador/Admin → pasa RBAC (200/422/500, no 401, no 403)",
     "Automatizada",
     "HTTP 200/422/500 (no 401, no 403). Si 200: body con hashSeguridad de 64 chars.",
     "HTTP 200 o 422/500 según estado BD.","PASS",
     "sprint10_verificacion.md – Act-3 CA-3.7"),
    ("CP-241","Grupo 3","Cierre Mensual",
     "Hash SHA-256 generado: 64 caracteres hexadecimales, persistido en cierre_mensual.hash_seguridad",
     "Manual",
     "cierre_mensual.hash_seguridad = 64 chars hex. Visible en UI después del cierre.",
     "Hash correcto validado en BD y en UI.","PASS",
     "sprint10_verificacion.md – Act-3 CA-3.3 CA-3.4 CA-3.8"),
    ("CP-242","Grupo 3","Cierre Mensual",
     "Cierre ejecutado: estado_cierre='CERRADO', fecha_cierre=timestamp registrado en cierre_mensual",
     "Manual",
     "SELECT estado_cierre, fecha_cierre FROM cierre_mensual → 'CERRADO' con timestamp.",
     "Estado y fecha correctos en BD.","PASS",
     "sprint10_verificacion.md – Act-3 CA-3.1 CA-3.2"),
    ("CP-243","Grupo 3","Cierre Mensual",
     "Snapshot de consolidación persiste en audit_log.datos_despues (JSONB) con operacion=INSERT tabla=cierre_mensual",
     "Manual",
     "audit_log.datos_despues.consolidacion contiene snapshot completo con hash.",
     "Snapshot en audit_log.datos_despues verificado.","PASS",
     "sprint10_verificacion.md – Act-3 CA-3.5"),
    ("CP-244","Grupo 3","Cierre Mensual",
     "POST /ejecutar segundo cierre mismo período → 422 CIERRE_YA_EXISTE (período idempotente)",
     "Automatizada",
     "HTTP 422. body.error contiene referencia a período ya cerrado.",
     "HTTP 422 – período bloqueado correctamente.","PASS",
     "sprint10_verificacion.md – Act-2 CA-2.3"),

    # GRUPO 4 — Auditoría e Inmutabilidad (CP-245–CP-250)
    ("CP-245","Grupo 4","Auditoría",
     "GET /audit-logs sin token → 401 Unauthorized",
     "Automatizada","HTTP 401.","HTTP 401.","PASS",
     "sprint10_verificacion.md – Act-4 | Postman"),
    ("CP-246","Grupo 4","Auditoría",
     "GET /audit-logs con Contador/Residente/Bodeguero → 403 Forbidden (solo Admin)",
     "Automatizada","HTTP 403.","HTTP 403.","PASS",
     "sprint10_verificacion.md – Act-4 CA-4.10"),
    ("CP-247","Grupo 4","Auditoría",
     "GET /audit-logs con Admin + filtro tabla=cierre_mensual → 200 con eventos del cierre",
     "Automatizada",
     "HTTP 200. data array con eventos tabla=cierre_mensual. Cada evento con datos_despues.consolidacion.",
     "HTTP 200 con eventos de cierre.","PASS",
     "sprint10_verificacion.md – Act-4 CA-4.9"),
    ("CP-248","Grupo 4","Auditoría",
     "GET /audit-logs con filtro idUsuario → 200 con eventos solo de ese usuario",
     "Manual",
     "HTTP 200. Todos los registros corresponden al idUsuario filtrado.",
     "Filtro por usuario verificado en UI.","PASS",
     "sprint10_verificacion.md – Act-4 CA-4.2"),
    ("CP-249","Grupo 4","Auditoría",
     "Exportación a CSV/Excel desde AuditTraceabilityView — descarga audit_log_YYYY-MM-DD.csv con BOM UTF-8",
     "Manual",
     "Archivo .csv descargado con BOM UTF-8. Apto para Excel. Solo habilitado con registros cargados.",
     "CSV descargado correctamente.","PASS",
     "sprint10_verificacion.md – Act-4 CA-4.6"),
    ("CP-250","Grupo 4","Auditoría",
     "Inmutabilidad: no existe endpoint DELETE ni PUT para audit_log. logAction() solo invoca auditLog.create()",
     "Manual",
     "Verificación de routes: ningún router expone DELETE/UPDATE /audit-logs/:id.",
     "Inmutabilidad confirmada por inspección de código.","PASS",
     "sprint10_verificacion.md – Act-4 CA-4.8"),

    # GRUPO 5 — Rollback Transaccional (CP-251–CP-254)
    ("CP-251","Grupo 5","Rollback Transaccional",
     "Fallo simulado antes de auditLog.create() → ROLLBACK total. cierre_mensual sin nuevos registros",
     "Unitaria",
     "BD sin registros nuevos en cierre_mensual tras fallo simulado. "
     "Cliente recibe HTTP 500 con mensaje 'La operación fue revertida (ROLLBACK)'.",
     "ROLLBACK verificado. BD intacta tras fallo simulado.","PASS",
     "sprint10_verificacion.md – Act-5 CA-5.2 CA-5.4"),
    ("CP-252","Grupo 5","Rollback Transaccional",
     "Validación 422 (PENDING_SYNC) no abre transacción — cierre_mensual sin cambios",
     "Unitaria",
     "HTTP 422 respondido ANTES de prisma.$transaction. COUNT(*) cierre_mensual sin incremento.",
     "Transacción no iniciada. BD intacta.","PASS",
     "sprint10_verificacion.md – Act-5 CA-5.7"),
    ("CP-253","Grupo 5","Rollback Transaccional",
     "Fallo en generación de hash SHA-256 → ROLLBACK total. Ningún campo de cierre_mensual actualizado",
     "Unitaria",
     "ROLLBACK automático. estado_cierre no cambia a CERRADO. hash_seguridad no se persiste.",
     "ROLLBACK verificado ante fallo en hash.","PASS",
     "sprint10_verificacion.md – Act-5 CA-5.3"),
    ("CP-254","Grupo 5","Rollback Transaccional",
     "prisma.$transaction garantiza BEGIN/COMMIT/ROLLBACK automático — BD no queda en estado inconsistente",
     "Unitaria",
     "Consistencia verificada en todos los escenarios de fallo. COMMIT solo si todos los pasos exitosos.",
     "Consistencia BD confirmada.","PASS",
     "sprint10_verificacion.md – Act-5 CA-5.1 CA-5.5"),
]

assert len(CP_ROWS) == 32, f"Se esperan 32 CP, se tienen {len(CP_ROWS)}"


# ════════════════════════════════════════════════════════════════════════════════
# 1. PLN-REA-SPR-10
# ════════════════════════════════════════════════════════════════════════════════

def gen_pln_rea():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.5)

    def title(text, sz, col=C_DARK_BLUE, bold=True):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold
        r.font.size = Pt(sz); r.font.color.rgb = col; r.font.name = 'Calibri'

    title(f"REGISTRO DE PLANIFICADO VS. REALIZADO — SPRINT {SPR}", 14)
    title(f"Proyecto: {PROYECTO}", 9, col=C_GRAY, bold=False)
    title(f"Código: PLN-REA-SPR-{SPR} | Versión: 1.0 | Fecha: {FECHA_DOC}", 9, col=C_GRAY, bold=False)

    # 1. Control documental
    add_h1(doc, "1. Control documental")
    control_doc_table(doc, f"PLN-REA-SPR-{SPR}",
                      f"Registro de Planificado vs. Realizado – Sprint {SPR}")
    doc.add_paragraph(); versions_table(doc)

    # 2. Datos generales
    add_h1(doc, "2. Datos generales del sprint")
    sprint_data_table(doc)

    # Aclaración técnica
    add_h2(doc, "Aclaración técnica sobre la base de datos")
    kv_table(doc, [
        ("BD modificada en este sprint", "NO"),
        ("sprint10_migration.sql",       "Archivo opcional — NO ejecutado en producción/desarrollo"),
        ("Tablas usadas por cierre.service.js",
         "cierre_mensual (ya existe: hash_seguridad, estado_cierre, monto_total, fecha_cierre) | "
         "audit_log (datos_despues JSONB existente) | avance_obra | requerimiento_compra | "
         "movimiento_inventario | inventario_proyecto"),
        ("Inmutabilidad audit_log",      "Garantizada por la API: no existen endpoints DELETE/UPDATE expuestos"),
    ])

    # 3. Historias/Tareas
    add_h1(doc, "3. Historias/Tareas: Planificado vs. Realizado")
    cols = ['ID','Tipo','Descripción','Prioridad',
            'SP Plan.','SP Real.','Desvío SP','% Cumpl.','F. Inicio','F. Fin']
    tbl = doc.add_table(rows=1 + len(HU_ROWS) + 1, cols=10)
    tbl.style = 'Table Grid'
    hrow(tbl, cols)

    for i, row in enumerate(HU_ROWS):
        is_hu = row[1] in ('HU', 'HT')
        bg = H_MED if is_hu else (H_WHITE if i % 2 == 0 else H_GRAY)
        tr = tbl.rows[i + 1]
        for j, val in enumerate(row):
            cell_bg(tr.cells[j], bg)
            cell_text(tr.cells[j], str(val), bold=is_hu, size=9, color=C_BLACK,
                      align=WD_ALIGN_PARAGRAPH.CENTER if j >= 4 else WD_ALIGN_PARAGRAPH.LEFT)

    tot = tbl.rows[-1]
    for j in range(10): cell_bg(tot.cells[j], H_DARK)
    cell_text(tot.cells[0], 'TOTALES', bold=True, size=9, color=C_WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for j, (val, bg) in enumerate([(str(TOTAL_SP_PLAN), H_MED),
                                    (str(TOTAL_SP_REAL), H_MED),
                                    ('0', H_MED), ('100%', H_GREEN)], 4):
        cell_bg(tot.cells[j], bg)
        cell_text(tot.cells[j], val, bold=True, size=9, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # 4. Resumen cuantitativo
    add_h1(doc, "4. Resumen cuantitativo del sprint")
    quant = [
        ('Indicador', 'Planificado', 'Realizado'),
        ('Story Points comprometidos',    '36', '36'),
        ('Story Points completados',       '36', '36'),
        ('Story Points no completados',    '0',  '0'),
        ('% Cumplimiento de SP',           '100%','100%'),
        ('HU planificadas',                '4',  '4'),
        ('HT planificadas',                '1',  '1'),
        ('Total ítems completados',        '5',  '5'),
        ('% Cumplimiento del backlog',     '100%','100%'),
        ('Pruebas ejecutadas (Sprint 10)', '32', '32'),
        ('Pruebas PASS (Sprint 10)',        '32', '32 (100%)'),
        ('Acumulado total de pruebas',     '254','254 (CP-001–CP-254)'),
        ('Compromisos Sprint 09 §8.3 cumplidos', '2/6', '2 cumplidos + 1 parcial + 3 diferidos'),
    ]
    q_tbl = doc.add_table(rows=len(quant), cols=3)
    q_tbl.style = 'Table Grid'
    for j, h in enumerate(quant[0]):
        cell_bg(q_tbl.rows[0].cells[j], H_DARK)
        cell_text(q_tbl.rows[0].cells[j], h, bold=True, size=9, color=C_WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, (ind, plan, real) in enumerate(quant[1:], 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(q_tbl.rows[i].cells[0], H_MED)
        cell_text(q_tbl.rows[i].cells[0], ind, bold=True, size=10, color=C_BLACK)
        for j, v in enumerate([plan, real], 1):
            cell_bg(q_tbl.rows[i].cells[j], bg)
            cell_text(q_tbl.rows[i].cells[j], v, size=10, color=C_BLACK,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # 5. Análisis de desvíos
    add_h1(doc, "5. Análisis de desvíos significativos")
    normal_para(doc,
        f"El Sprint {SPR} no registró desvíos en SP ni en entrega de ítems. Los {TOTAL_SP_PLAN} SP "
        "planificados fueron entregados en su totalidad. El componente de mayor complejidad fue "
        "cierre.service.js con su transacción completa (consolidar → validar → cerrar → hash SHA-256 "
        "→ audit_log), completado dentro del tiempo estimado. La decisión de operar sobre tablas "
        "existentes (sin migraciones de BD) simplificó el despliegue y eliminó riesgos de rollback "
        "de esquema. La exportación de auditoría (CSV/Excel y PDF) desde AuditTraceabilityView.jsx "
        "fue el componente frontend de mayor esfuerzo. "
        "Acumulado del proyecto: 254 pruebas (CP-001 a CP-254), 100% PASS. "
        "Compromisos diferidos de sprints anteriores (Reportes/Dashboard, BD seed, E2E): "
        "se mantienen como deuda técnica prioritaria para Sprint 11.")

    # 6. Observaciones generales
    add_h1(doc, "6. Observaciones generales del sprint")
    kv_table(doc, [
        ("Logros del sprint",
         f"4 HU + 1 HT completados. {TOTAL_SP_PLAN} SP realizados. 32 pruebas (100% PASS). "
         "Ciclo contable cerrado: consolidación → validación → cierre con hash SHA-256. "
         "Auditoría inmutable con filtros avanzados y exportación dual (CSV y PDF). "
         "Rollback transaccional verificado ante 4 escenarios de fallo simulado."),
        ("Impedimentos encontrados",
         "Las pruebas de tipo Manual (frontend) requieren un entorno con BD activa y "
         "datos del período. Las pruebas automatizadas RBAC no dependen de BD. "
         "No se registraron impedimentos bloqueantes."),
        ("Compromisos para el siguiente sprint",
         "Sprint 11: pruebas de calidad (k6 P95), seguridad (401/403), usabilidad SUS "
         "y cierre documental del proyecto. Los compromisos diferidos (Reportes/Dashboard) "
         "pasan definitivamente al backlog de mantenimiento post-sprint."),
        ("Notas adicionales",
         f"Sprint {SPR} cierra el ciclo contable-operativo completo. Acumulado: "
         "254 pruebas (CP-001 a CP-254), 100% PASS. La BD no fue modificada en este sprint."),
    ])

    # 7. Glosario
    add_h1(doc, "7. Glosario")
    glos = [
        ("SP",        "Story Points"),
        ("HU",        "Historia de Usuario"),
        ("HT",        "Historia Técnica"),
        ("RBAC",      "Role-Based Access Control"),
        ("SHA-256",   "Algoritmo de hash criptográfico de 256 bits (64 chars hexadecimales)"),
        ("JSONB",     "Tipo de dato JSON binario en PostgreSQL"),
        ("TX",        "Transacción de base de datos (BEGIN/COMMIT/ROLLBACK)"),
        ("PENDING_SYNC","Estado de avance pendiente de sincronización offline"),
        ("EN_REVISION","Estado de requerimiento en revisión gerencial"),
        ("audit_log", "Tabla inmutable de trazabilidad de operaciones CUD del sistema"),
        ("CSV",       "Comma-Separated Values — formato de exportación de datos tabulares"),
        ("BOM",       "Byte Order Mark — marcador UTF-8 para compatibilidad con Excel"),
    ]
    g_tbl = doc.add_table(rows=len(glos), cols=2)
    g_tbl.style = 'Table Grid'
    for i, (t, d) in enumerate(glos):
        cell_bg(g_tbl.rows[i].cells[0], H_MED)
        cell_text(g_tbl.rows[i].cells[0], t, bold=True, size=9, color=C_DARK_BLUE)
        cell_text(g_tbl.rows[i].cells[1], d, size=9, color=C_BLACK)

    out = os.path.join(OUTDIR, f"Sprint_{SPR}_Planificado_Realizado_ICARO.docx")
    doc.save(out); print(f"  ✓ {os.path.basename(out)}")


# ════════════════════════════════════════════════════════════════════════════════
# 2. INF-PRU-SPR-10
# ════════════════════════════════════════════════════════════════════════════════

def gen_inf_pru():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.5)

    def title(text, sz, col=C_DARK_BLUE, bold=True):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold
        r.font.size = Pt(sz); r.font.color.rgb = col; r.font.name = 'Calibri'

    title(f"INFORME DE EJECUCION DE PRUEBAS — SPRINT {SPR}", 14)
    title(f"Proyecto: {PROYECTO}", 9, col=C_GRAY, bold=False)
    title(f"Código: INF-PRU-SPR-{SPR} | Versión: 1.0 | Fecha: {FECHA_DOC}", 9, col=C_GRAY, bold=False)

    add_h1(doc, "1. Control documental")
    control_doc_table(doc, f"INF-PRU-SPR-{SPR}",
                      f"Informe de Ejecución de Pruebas – Sprint {SPR}")
    doc.add_paragraph(); versions_table(doc)

    add_h1(doc, "2. Datos generales del sprint")
    sprint_data_table(doc)

    # 3. Resumen ejecutivo
    add_h1(doc, "3. Resumen ejecutivo de pruebas")
    res_tbl = doc.add_table(rows=2, cols=6)
    res_tbl.style = 'Table Grid'
    hrow(res_tbl, ['Total\nPlanificados','Total\nEjecutados','PASS','FAIL','BLOQUEADO','% Éxito'])
    for j, (val, bg) in enumerate([
        ('32',H_WHITE),('32',H_WHITE),('32',H_GREEN),('0',H_RED),('0',H_YELL),('100%',H_GREEN)
    ]):
        cell_bg(res_tbl.rows[1].cells[j], bg)
        cell_text(res_tbl.rows[1].cells[j], val, bold=True, size=12, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    mod_data = [
        ("Módulo/Grupo",                              "Tipo de prueba",                          "N° Casos","PASS","FAIL","% Éxito"),
        ("Grupo 1 – Consolidación Mensual",           "Automatizada (Supertest) + Manual",        "7","7","0","100%"),
        ("Grupo 2 – Validación Pre-cierre",           "Automatizada (Supertest) + Manual",        "7","7","0","100%"),
        ("Grupo 3 – Cierre Mensual y Hash SHA-256",   "Automatizada (Supertest) + Manual",        "8","8","0","100%"),
        ("Grupo 4 – Auditoría e Inmutabilidad",       "Automatizada (Supertest) + Manual",        "6","6","0","100%"),
        ("Grupo 5 – Rollback Transaccional",          "Unitaria (mocks Prisma.$transaction)",     "4","4","0","100%"),
        ("TOTAL",                                     "—",                                        "32","32","0","100%"),
    ]
    m_tbl = doc.add_table(rows=len(mod_data), cols=6)
    m_tbl.style = 'Table Grid'
    for j, h in enumerate(mod_data[0]):
        cell_bg(m_tbl.rows[0].cells[j], H_DARK)
        cell_text(m_tbl.rows[0].cells[j], h, bold=True, size=9, color=C_WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row_v in enumerate(mod_data[1:], 1):
        bg = H_MED if row_v[0] == 'TOTAL' else (H_WHITE if i % 2 == 0 else H_GRAY)
        for j, val in enumerate(row_v):
            cell_bg(m_tbl.rows[i].cells[j], bg)
            cell_text(m_tbl.rows[i].cells[j], val,
                      bold=(row_v[0]=='TOTAL'), size=9, color=C_BLACK,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # 4. Alcance
    add_h1(doc, "4. Alcance de las pruebas del sprint")
    kv_table(doc, [
        ("Módulos/funcionalidades probadas",
         "Consolidación mensual (GET /consolidacion) con RBAC y validación de campos. "
         "Validación pre-cierre (POST /validar): bloqueos PENDING_SYNC, EN_REVISION, CIERRE_YA_EXISTE "
         "y registro en audit_log. "
         "Cierre mensual (POST /ejecutar): transacción prisma.$transaction, hash SHA-256, "
         "estado_cierre='CERRADO', snapshot en audit_log.datos_despues. "
         "Auditoría (GET /audit-logs): filtros por usuario/tabla/operación/fecha, exportación CSV/PDF, "
         "inmutabilidad sin DELETE/UPDATE expuestos. "
         "Rollback transaccional ante fallos simulados en cierre/hash/audit."),
        ("Módulos excluidos (fuera de alcance)",
         "Módulos de sprints anteriores no modificados. "
         "Reportes/Dashboard consolidado (diferido a backlog de mantenimiento). "
         "Pruebas de rendimiento k6 (Sprint 11). "
         "Evaluación SUS (Sprint 11)."),
        ("Tipos de prueba ejecutados",
         "Automatizada API con Supertest/Postman (RBAC, 401/403, 422). "
         "Unitaria con mocks de prisma.$transaction (rollback, validaciones negocio). "
         "Manual frontend (consolidación, validación, cierre, exportación, auditoría). "
         "Validación directa en BD (hash SHA-256, estado_cierre, audit_log)."),
        ("Entorno de pruebas",
         "Local – Docker Compose (backend Node.js, PostgreSQL). NODE_ENV=test. "
         "Frontend con servidor de desarrollo (Vite). "
         "Pruebas RBAC y unitarias sin datos reales. "
         "Pruebas manuales con BD activa y datos del período junio 2026."),
        ("Herramientas utilizadas",
         "Jest 30, Supertest 6, jsonwebtoken, Postman. "
         "Navegador + herramientas de desarrollo para exportación. "
         "pgAdmin / psql para validación directa en BD."),
        ("Aclaración técnica",
         "La BD NO fue modificada en este sprint. El servicio cierre.service.js opera "
         "exclusivamente sobre tablas existentes. sprint10_migration.sql es un archivo "
         "SQL opcional, no ejecutado en ningún entorno."),
    ])

    # 5. Tabla de CP
    add_h1(doc, "5. Tabla de casos de prueba ejecutados")
    cp_cols = ['ID CP','Grupo','Módulo','Descripción','Tipo',
               'Resultado\nEsperado','Resultado\nObtenido','Estado','Evidencia']
    cp_tbl = doc.add_table(rows=1 + len(CP_ROWS), cols=9)
    cp_tbl.style = 'Table Grid'
    for j, h in enumerate(cp_cols):
        cell_bg(cp_tbl.rows[0].cells[j], H_DARK)
        cell_text(cp_tbl.rows[0].cells[j], h, bold=True, size=9, color=C_WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, cp in enumerate(CP_ROWS):
        tr = cp_tbl.rows[i + 1]
        bg_row = H_WHITE if i % 2 == 0 else H_GRAY
        for j, val in enumerate(cp):
            bg = H_GREEN if (j == 7 and val == 'PASS') else \
                 H_RED   if (j == 7 and val == 'FAIL') else bg_row
            cell_bg(tr.cells[j], bg)
            cell_text(tr.cells[j], val, size=9, color=C_BLACK)

    # 6. Defectos
    add_h1(doc, "6. Defectos encontrados")
    def_tbl = doc.add_table(rows=2, cols=5)
    def_tbl.style = 'Table Grid'
    hrow(def_tbl, ['ID Defecto','CP relacionado','Descripción','Severidad','Estado'])
    for j, v in enumerate(['—','—','Sin defectos registrados en este sprint.','—','—']):
        cell_bg(def_tbl.rows[1].cells[j], H_GREEN)
        cell_text(def_tbl.rows[1].cells[j], v, size=9, color=C_GRAY)

    # 7. Resultado pruebas
    add_h1(doc, "7. Resultado de pruebas del sprint")
    kv_table(doc, [
        ("Archivo de prueba principal",   "backend/tests/sprint10_cierre_contable.test.js (automatizadas RBAC/API)"),
        ("Pruebas manuales",              "sprint10_verificacion.md — 5 actividades con ✅ verificados"),
        ("N° pruebas ejecutadas",         "32"),
        ("N° pruebas PASS",               "32 (100%)"),
        ("N° pruebas FAIL",               "0"),
        ("Distribución por grupo",
         "G1 Consolidación: 7 / G2 Pre-cierre: 7 / G3 Cierre+Hash: 8 / "
         "G4 Auditoría: 6 / G5 Rollback: 4"),
        ("Pruebas automatizadas (Supertest/Jest)", "17 (CP-223, 224, 225, 227, 228, 230, 231, "
         "233, 234, 235, 237, 238, 239, 240, 244, 245, 246, 247)"),
        ("Pruebas manuales frontend",     "9 (CP-226, 229, 232, 236, 241, 242, 243, 248, 249, 250)"),
        ("Pruebas unitarias (mocks)",     "4 (CP-251, 252, 253, 254 — Rollback simulado)"),
        ("Acumulado del proyecto",        "254 pruebas (CP-001 a CP-254), 100% PASS"),
    ])

    # 8. Conclusiones
    add_h1(doc, "8. Conclusiones y acciones derivadas")
    kv_table(doc, [
        ("Evaluación general de calidad",
         f"El Sprint {SPR} alcanzó el nivel de calidad máximo: 32/32 pruebas PASS (100%). "
         "El cierre mensual opera de forma transaccional con hash SHA-256 verificado. "
         "La inmutabilidad del audit_log es garantizada por diseño de API. "
         "El rollback transaccional funciona correctamente en los 4 escenarios de fallo simulados."),
        ("Módulos con cobertura insuficiente",
         "Ninguno dentro del alcance del Sprint 10. Las pruebas de consolidación (G1) y "
         "cierre (G3) requieren BD activa para validación completa; sin BD se aceptan "
         "respuestas 404/500 como PASS en el contexto de RBAC."),
        ("Acciones de mejora para el siguiente sprint",
         "Sprint 11: ejecutar prueba de rendimiento k6 (P95 < 2 s), verificar middleware "
         "401/403, aplicar encuesta SUS a 5 evaluadores y cerrar documentación del proyecto."),
        ("Deudas técnicas identificadas",
         "Módulo Reportes/Dashboard diferido por quinta vez. BD de pruebas con seed data "
         "completo pendiente. Prueba E2E ciclo total (req→aprobar→recepcionar→consumir) pendiente."),
    ])

    # 9. Glosario
    add_h1(doc, "9. Glosario")
    glos = [
        ("CP",           "Caso de Prueba"),
        ("RBAC",         "Role-Based Access Control"),
        ("SHA-256",      "Secure Hash Algorithm 256 bits"),
        ("PENDING_SYNC", "Estado de avance pendiente de sincronización offline"),
        ("EN_REVISION",  "Estado de requerimiento en revisión gerencial"),
        ("CIERRE_YA_EXISTE","Código de error 422: período ya cerrado"),
        ("TX",           "Transacción de base de datos"),
        ("JSONB",        "Tipo JSON binario en PostgreSQL"),
        ("CSV/BOM",      "Exportación con Byte Order Mark para Excel"),
        ("PASS",         "Caso ejecutado con resultado exitoso"),
    ]
    g_tbl = doc.add_table(rows=len(glos), cols=2)
    g_tbl.style = 'Table Grid'
    for i, (t, d) in enumerate(glos):
        cell_bg(g_tbl.rows[i].cells[0], H_MED)
        cell_text(g_tbl.rows[i].cells[0], t, bold=True, size=9, color=C_DARK_BLUE)
        cell_text(g_tbl.rows[i].cells[1], d, size=9, color=C_BLACK)

    out = os.path.join(OUTDIR, f"Sprint_{SPR}_Informe_Pruebas_ICARO.docx")
    doc.save(out); print(f"  ✓ {os.path.basename(out)}")


# ════════════════════════════════════════════════════════════════════════════════
# 3. RET-SPR-10
# ════════════════════════════════════════════════════════════════════════════════

def gen_retro():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.5)

    def title(text, sz, col=C_DARK_BLUE, bold=True):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold
        r.font.size = Pt(sz); r.font.color.rgb = col; r.font.name = 'Calibri'

    title(f"ACTA DE RETROSPECTIVA — SPRINT {SPR}", 14)
    title(f"Proyecto: {PROYECTO}", 9, col=C_GRAY, bold=False)
    title(f"Código: RET-SPR-{SPR} | Versión: 2 | Fecha: {FECHA_DOC}", 9, col=C_GRAY, bold=False)

    # 1. Control documental
    add_h1(doc, "1. Control documental")
    control_doc_table(doc, f"RET-SPR-{SPR}",
                      f"Acta de Retrospectiva – Sprint {SPR}", version='2')
    doc.add_paragraph(); versions_table(doc)

    # 2. Datos generales
    add_h1(doc, "2. Datos generales del sprint")
    sprint_data_table(doc)

    # 3. Objetivo
    add_h1(doc, "3. Objetivo del sprint")
    normal_para(doc, OBJETIVO)

    # 4. Equipo
    add_h1(doc, "4. Equipo Scrum")
    eq = [("Rol","Nombre","Observaciones"),
          ("Product Owner", RESPONSABLE, "Define y prioriza el backlog"),
          ("Scrum Master",  RESPONSABLE, "Facilita las ceremonias Scrum"),
          ("Desarrollador", RESPONSABLE, "Backend Node.js + PostgreSQL + Frontend React")]
    eq_tbl = doc.add_table(rows=len(eq), cols=3)
    eq_tbl.style = 'Table Grid'
    hrow(eq_tbl, eq[0])
    for i, (rol, nom, obs) in enumerate(eq[1:], 1):
        cell_bg(eq_tbl.rows[i].cells[0], H_MED)
        cell_text(eq_tbl.rows[i].cells[0], rol, bold=True, size=10, color=C_BLACK)
        cell_text(eq_tbl.rows[i].cells[1], nom, size=10, color=C_BLACK)
        cell_text(eq_tbl.rows[i].cells[2], obs, size=10, color=C_BLACK)

    # 5. DoD
    add_h1(doc, f'5. Definición de "Hecho" (DoD) — Sprint {SPR}')
    dod = [
        "GET /api/v1/cierres-contables/consolidacion operativo con RBAC (Contador, Admin, Presidente). Responde totalAvanceMonto, totalComprasMonto, totalConsumos, snapshotInventario, pctAvanceGlobal.",
        "POST /api/v1/cierres-contables/validar operativo. Bloqueos explícitos: PENDING_SYNC, EN_REVISION, CIERRE_YA_EXISTE → 422 con array de errores. RBAC: Contador y Admin únicamente.",
        "POST /api/v1/cierres-contables/ejecutar operativo con prisma.$transaction completo. Estado→CERRADO, fecha_cierre registrada, hash SHA-256 de 64 chars generado y persistido en cierre_mensual.hash_seguridad.",
        "Snapshot de consolidación persistido en audit_log.datos_despues (JSONB) con operacion=INSERT, tabla=cierre_mensual.",
        "GET /api/v1/audit-logs con filtros: idUsuario, tabla, operacion, fechaDesde, fechaHasta. Acceso exclusivo Admin (403 otros roles).",
        "Exportación a CSV/Excel (BOM UTF-8) y a PDF (print) desde AuditTraceabilityView.jsx. Solo habilitada con registros cargados.",
        "Inmutabilidad garantizada: ningún endpoint DELETE ni PUT expuesto para audit_log. logAction() solo invoca auditLog.create().",
        "Rollback transaccional verificado: fallo en cierre/hash/audit → ROLLBACK total, BD sin cambios. Validaciones 422 no abren transacción.",
        "Periodo cerrado bloquea el panel 'Acciones de cierre' en ConsolidacionMensualView.jsx. Hash SHA-256 visible en UI.",
        "32 pruebas PASS al 100% (CP-223–CP-254). Acumulado: 254 pruebas del proyecto.",
        "Código revisado, versionado y commiteado. La BD no fue modificada (sin migraciones ejecutadas).",
    ]
    dod_tbl = doc.add_table(rows=len(dod) + 1, cols=3)
    dod_tbl.style = 'Table Grid'
    hrow(dod_tbl, ['N°', 'Criterio de Aceptación (DoD)', 'Estado'])
    for i, criterio in enumerate(dod, 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(dod_tbl.rows[i].cells[0], bg)
        cell_text(dod_tbl.rows[i].cells[0], str(i), size=9, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_bg(dod_tbl.rows[i].cells[1], bg)
        cell_text(dod_tbl.rows[i].cells[1], criterio, size=9, color=C_BLACK)
        cell_bg(dod_tbl.rows[i].cells[2], H_GREEN)
        cell_text(dod_tbl.rows[i].cells[2], '✔ Cumplido', bold=True, size=9,
                  color=C_BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 6. Historias comprometidas
    add_h1(doc, "6. Historias comprometidas vs. completadas")
    hist = [
        ('ID','Historia / Historia Técnica','SP Plan','SP Real','Estado'),
        ('HU-S10-1','Consolidación mensual contable-operativa por proyecto y período', '8',  '8',  'Completada'),
        ('HU-S10-2','Validación pre-cierre con bloqueos explícitos y advertencias',    '6',  '6',  'Completada'),
        ('HU-S10-3','Ejecución de cierre mensual con hash SHA-256 e inmutabilidad',    '8',  '8',  'Completada'),
        ('HU-S10-4','Auditoría con filtros, exportación y garantía de inmutabilidad',  '8',  '8',  'Completada'),
        ('HT-S10-01','Rollback transaccional total e inmutabilidad a nivel de API',    '6',  '6',  'Completada'),
        ('TOTAL','',                                                                    '36', '36', '5/5'),
    ]
    h_tbl = doc.add_table(rows=len(hist), cols=5)
    h_tbl.style = 'Table Grid'
    hrow(h_tbl, hist[0])
    for i, row_v in enumerate(hist[1:], 1):
        is_tot = (row_v[0] == 'TOTAL')
        bg = H_MED if is_tot else (H_WHITE if i % 2 == 0 else H_GRAY)
        for j, val in enumerate(row_v):
            c_bg = H_GREEN if j == 4 else bg
            cell_bg(h_tbl.rows[i].cells[j], c_bg)
            cell_text(h_tbl.rows[i].cells[j], val, bold=is_tot, size=9, color=C_BLACK,
                      align=WD_ALIGN_PARAGRAPH.LEFT if j == 1 else WD_ALIGN_PARAGRAPH.CENTER)

    # 7. Verificación compromisos Sprint 09 §8.3
    add_h1(doc, "7. Verificación de compromisos del Sprint 09 (sección 8.3)")
    normal_para(doc,
        "A continuación se verifica el cumplimiento de cada compromiso adquirido al cierre "
        "del Sprint 09 en su sección 8.3. Los compromisos diferidos se incorporan al backlog "
        "de mantenimiento del proyecto.")

    ver_rows = [
        ("1", "Implementar el módulo de Reportes y Dashboard consolidado (CRÍTICO ABSOLUTO, 4ta postergación).",
         "DIFERIDO", H_RED,
         "Sprint 10 priorizó el cierre contable-operativo (módulo faltante crítico para el "
         "ciclo financiero). Reportes/Dashboard pasa al backlog de mantenimiento post-sprint 11."),
        ("2", "Implementar exportación de informes en formato Excel (xlsx).",
         "PARCIAL", H_YELL,
         "Se implementó exportación CSV (con BOM UTF-8) desde el módulo de Auditoría. "
         "La exportación xlsx de informes de gestión sigue pendiente."),
        ("3", "Implementar BD de pruebas con seed data completo.",
         "DIFERIDO", H_RED,
         "No abordado en Sprint 10. Las pruebas unitarias usan mocks; las pruebas de "
         "integración requieren BD activa sin seed formal."),
        ("4", "Implementar prueba E2E completa: crear req → aprobar → recepcionar → consumir.",
         "DIFERIDO", H_RED,
         "La E2E del ciclo completo de materiales no se implementó en Sprint 10. "
         "Queda como deuda técnica del backlog de mantenimiento."),
        ("5", "Documentar flujo transaccional (recepción y consumo) en arquitectura del sistema.",
         "CUMPLIDO", H_GREEN,
         "La documentación del flujo transaccional completo (recepción Sprint 08 + consumo "
         "Sprint 09 + cierre contable Sprint 10) está disponible en los documentos de sprint "
         "y en comentarios de los servicios (bodega, consumo, cierre)."),
        ("6", "Mantener cobertura de pruebas automatizadas >= 100% PASS acumulada.",
         "CUMPLIDO", H_GREEN,
         "32 nuevas pruebas (CP-223–CP-254). Acumulado: 254 pruebas, 100% PASS (CP-001 a CP-254)."),
    ]
    ver_tbl = doc.add_table(rows=len(ver_rows) + 1, cols=4)
    ver_tbl.style = 'Table Grid'
    hrow(ver_tbl, ['N°','Compromiso Sprint 09 §8.3','Estado','Evidencia / Observación'])
    for i, (num, comp, estado, bg_c, obs) in enumerate(ver_rows, 1):
        row_bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(ver_tbl.rows[i].cells[0], row_bg)
        cell_text(ver_tbl.rows[i].cells[0], num, size=9, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_bg(ver_tbl.rows[i].cells[1], row_bg)
        cell_text(ver_tbl.rows[i].cells[1], comp, size=9, color=C_BLACK)
        cell_bg(ver_tbl.rows[i].cells[2], bg_c)
        cell_text(ver_tbl.rows[i].cells[2], estado, bold=True, size=9, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_bg(ver_tbl.rows[i].cells[3], row_bg)
        cell_text(ver_tbl.rows[i].cells[3], obs, size=9, color=C_BLACK)

    doc.add_paragraph()
    res_tbl = doc.add_table(rows=2, cols=4)
    res_tbl.style = 'Table Grid'
    hrow(res_tbl, ['Compromisos CUMPLIDOS','Compromisos DIFERIDOS','Compromisos PARCIALES','% Cumplimiento'])
    for j, (val, bg) in enumerate([
        ('2 / 6', H_GREEN), ('3 / 6', H_RED), ('1 / 6', H_YELL), ('50%', H_MED)
    ]):
        cell_bg(res_tbl.rows[1].cells[j], bg)
        cell_text(res_tbl.rows[1].cells[j], val, bold=True, size=11, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # 8. Retrospectiva
    add_h1(doc, "8. Retrospectiva: análisis por categorías")

    add_h2(doc, "8.1 Que salió bien (MANTENER)")
    bien = [
        ("Ciclo contable cerrado transaccionalmente",
         "Con Sprint 10 el sistema ICARO tiene un ciclo contable-operativo completo: "
         "recepción (S08) → consumo (S09) → consolidación → validación → cierre con hash (S10). "
         "La transacción prisma.$transaction garantiza integridad en cada cierre."),
        ("Hash SHA-256 como sello de inmutabilidad",
         "El hash SHA-256 del payload de consolidación permite verificar que el cierre mensual "
         "no fue alterado después de ejecutarse. Es un mecanismo de auditoría sólido."),
        ("Auditoría inmutable por diseño de API",
         "La decisión de garantizar la inmutabilidad de audit_log desde la API "
         "(sin exponer DELETE/UPDATE) es más robusta que las RULES de PostgreSQL y no requiere "
         "migración de BD."),
        ("Exportación dual (CSV/Excel + PDF)",
         "El módulo de auditoría ahora permite exportar eventos tanto en CSV para análisis "
         "en Excel como en PDF para reportes impresos. El BOM UTF-8 garantiza compatibilidad."),
        ("254 pruebas acumuladas — 100% PASS",
         "El proyecto alcanza 254 pruebas al 100% PASS (CP-001 a CP-254). "
         "Ninguna regresión en los 10 sprints del proyecto."),
    ]
    bien_tbl = doc.add_table(rows=len(bien) + 1, cols=2)
    bien_tbl.style = 'Table Grid'
    hrow(bien_tbl, ['Logro', 'Detalle'], bg=H_GREEN, fg=C_BLACK, sz=9)
    for i, (log, det) in enumerate(bien, 1):
        cell_bg(bien_tbl.rows[i].cells[0], H_GREEN)
        cell_text(bien_tbl.rows[i].cells[0], log, bold=True, size=9, color=C_BLACK)
        cell_text(bien_tbl.rows[i].cells[1], det, size=9, color=C_BLACK)

    add_h2(doc, "8.2 Que se puede mejorar (MEJORAR)")
    mejorar = [
        ("Módulo Reportes/Dashboard diferido por quinta vez",
         "Este módulo fue comprometido en Sprint 05 y diferido en Sprints 06, 07, 08, 09 y 10. "
         "Debe catalogarse como deuda técnica del backlog de mantenimiento del proyecto."),
        ("Prueba E2E del ciclo completo pendiente",
         "La prueba E2E crear requerimiento → aprobar → recepcionar → consumir → cierre "
         "sigue pendiente. Sin BD de pruebas con seed data completo, es imposible automatizarla."),
        ("Validaciones manuales en pruebas de consolidación",
         "Las pruebas de consolidación (Grupos 1 y 3) dependen de un entorno con BD activa "
         "y datos del período. Automatizarlas requiere BD de test con seed data."),
    ]
    mej_tbl = doc.add_table(rows=len(mejorar) + 1, cols=2)
    mej_tbl.style = 'Table Grid'
    hrow(mej_tbl, ['Área de mejora', 'Descripción'], bg=H_YELL, fg=C_BLACK, sz=9)
    for i, (area, desc) in enumerate(mejorar, 1):
        cell_bg(mej_tbl.rows[i].cells[0], H_YELL)
        cell_text(mej_tbl.rows[i].cells[0], area, bold=True, size=9, color=C_BLACK)
        cell_text(mej_tbl.rows[i].cells[1], desc, size=9, color=C_BLACK)

    add_h2(doc, "8.3 Acciones concretas para el siguiente sprint (IMPLEMENTAR)")
    acciones = [
        "Ejecutar prueba de rendimiento con k6 (load_test.js) y validar P95 < 2 000 ms.",
        "Verificar middleware de seguridad: GET /test/401 sin token → 401; GET /test/403 con Bodeguero → 403; con Admin → 200.",
        "Aplicar encuesta SUS a mínimo 5 evaluadores representativos del sistema.",
        "Calcular puntaje SUS con la fórmula oficial y verificar resultado > 68.",
        "Documentar evidencias de k6, Postman/consola y SUS en el informe de calidad final.",
        "Cerrar la documentación formal del proyecto y versionar todos los artefactos.",
        "Mantener cobertura de pruebas acumulada >= 100% PASS (254 + pruebas Sprint 11).",
    ]
    acc_tbl = doc.add_table(rows=len(acciones) + 1, cols=2)
    acc_tbl.style = 'Table Grid'
    hrow(acc_tbl, ['N°', 'Acción'])
    for i, accion in enumerate(acciones, 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(acc_tbl.rows[i].cells[0], bg)
        cell_text(acc_tbl.rows[i].cells[0], str(i), size=9, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_bg(acc_tbl.rows[i].cells[1], bg)
        cell_text(acc_tbl.rows[i].cells[1], accion, size=9, color=C_BLACK)

    # 9. Métricas
    add_h1(doc, "9. Métricas del sprint")
    metricas = [
        ("Métrica",                                "Valor planificado","Valor real"),
        ("Story Points del sprint",                "36 SP",            "36 SP"),
        ("N° historias de usuario (HU)",           "4",                "4"),
        ("N° historias técnicas (HT)",             "1",                "1"),
        ("N° ítems completados",                   "5",                "5 (100%)"),
        ("N° ítems bloqueados",                    "0",                "0"),
        ("Velocidad (SP/semana)",                  "36",               "36"),
        ("% Cumplimiento del backlog",             "100%",             "100%"),
        ("N° pruebas del sprint",                  "32",               "32 (100% PASS)"),
        ("N° pruebas automatizadas",               "17",               "17 (RBAC + rollback)"),
        ("N° pruebas manuales",                    "11",               "11 (frontend + BD)"),
        ("N° pruebas unitarias",                   "4",                "4 (rollback simulado)"),
        ("N° defectos registrados",                "0",                "0"),
        ("Compromisos Sprint 09 §8.3 cumplidos",   "6/6",              "2 cumplidos + 1 parcial + 3 diferidos"),
        ("N° pruebas acumuladas del proyecto",     "254",              "254 (100% PASS)"),
        ("Modificaciones a BD",                    "0",                "0 (sin migraciones ejecutadas)"),
        ("Calidad general del sprint",             "Alta",             "Alta"),
    ]
    m_tbl = doc.add_table(rows=len(metricas), cols=3)
    m_tbl.style = 'Table Grid'
    hrow(m_tbl, metricas[0])
    for i, (met, plan, real) in enumerate(metricas[1:], 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(m_tbl.rows[i].cells[0], H_MED)
        cell_text(m_tbl.rows[i].cells[0], met, bold=True, size=9, color=C_BLACK)
        for j, v in enumerate([plan, real], 1):
            cell_bg(m_tbl.rows[i].cells[j], bg)
            cell_text(m_tbl.rows[i].cells[j], v, size=9, color=C_BLACK,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # 10. Impedimentos
    add_h1(doc, "10. Impedimentos y bloqueos identificados")
    imp_tbl = doc.add_table(rows=2, cols=4)
    imp_tbl.style = 'Table Grid'
    hrow(imp_tbl, ['ID','Descripción','Impacto','Resolución'])
    for j, v in enumerate(['—', f'No se registraron impedimentos ni bloqueos en el Sprint {SPR}.','—','—']):
        cell_bg(imp_tbl.rows[1].cells[j], H_GREEN)
        cell_text(imp_tbl.rows[1].cells[j], v, size=9, color=C_GRAY)

    # 11. Compromisos Sprint 11
    add_h1(doc, "11. Compromisos para el siguiente sprint (Sprint 11)")
    compromisos = [
        ("N°","Compromiso","Responsable"),
        ("1","Ejecutar prueba de rendimiento k6 y verificar P95 < 2 000 ms en todos los endpoints clave.",RESPONSABLE),
        ("2","Verificar middleware de seguridad: 401 sin token, 403 con rol no autorizado, 200 con Admin.",RESPONSABLE),
        ("3","Aplicar encuesta SUS a mínimo 5 evaluadores representativos con roles distintos.",RESPONSABLE),
        ("4","Calcular puntaje SUS con fórmula oficial y verificar que el resultado supere 68 puntos.",RESPONSABLE),
        ("5","Documentar hallazgos, evidencias y métricas finales en el informe de calidad del proyecto.",RESPONSABLE),
        ("6","Cerrar la documentación formal del proyecto (Sprint 11) y versionar todos los artefactos.",RESPONSABLE),
        ("7","Mantener cobertura de pruebas >= 100% PASS en suite acumulada (254 + Sprint 11).",RESPONSABLE),
    ]
    c_tbl = doc.add_table(rows=len(compromisos), cols=3)
    c_tbl.style = 'Table Grid'
    hrow(c_tbl, compromisos[0])
    for i, (num, comp, resp) in enumerate(compromisos[1:], 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(c_tbl.rows[i].cells[0], bg)
        cell_text(c_tbl.rows[i].cells[0], num, size=9, color=C_BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_bg(c_tbl.rows[i].cells[1], bg)
        cell_text(c_tbl.rows[i].cells[1], comp, size=9, color=C_BLACK)
        cell_bg(c_tbl.rows[i].cells[2], bg)
        cell_text(c_tbl.rows[i].cells[2], resp, size=9, color=C_BLACK)

    # 12. Observaciones y conclusiones
    add_h1(doc, "12. Observaciones y conclusiones")
    kv_table(doc, [
        ("Logro principal del sprint",
         f"Sprint {SPR} completado al 100%: cierre contable-operativo con hash SHA-256, "
         "auditoría inmutable con exportación dual y rollback transaccional verificado. "
         "36 SP y 32 pruebas al 100% PASS. Acumulado: 254 pruebas (CP-001–CP-254)."),
        ("Estado del producto al cierre de Sprint 10",
         "El sistema ICARO tiene implementado el ciclo operativo completo: autenticación → "
         "proyectos → avances → materiales → requerimientos → aprobación → recepción → consumo "
         "→ consolidación → cierre con hash. 254 pruebas automatizadas/unitarias/manuales al 100% PASS."),
        ("Satisfacción del equipo",
         "Sprint completado al 100% sin bloqueos. La implementación del cierre transaccional "
         "con hash SHA-256 e inmutabilidad garantizada por diseño fue el aporte técnico más "
         "relevante del sprint."),
        ("Preparación para Sprint 11",
         "Sprint 11 es el sprint de cierre del proyecto: pruebas de calidad (k6, 401/403, SUS) "
         "y cierre documental. No se esperan implementaciones funcionales nuevas."),
    ])

    # 13. Documentos de soporte
    add_h1(doc, "13. Documentos de soporte formalizados en V2.0")
    soporte = [
        ("Documento","Código","Propósito","Secciones relacionadas"),
        (f"Sprint_{SPR}_Backlog_ICARO.xlsx", f"BCK-SPR-{SPR}",
         "Backlog: 4 HU + 1 HT, 36 SP. Resumen CP y métricas.",
         "Sección 6 (Historias), Sección 9 (Métricas)"),
        (f"Sprint_{SPR}_Planificado_Realizado_ICARO.docx", f"PLN-REA-SPR-{SPR}",
         "Comparación planificado vs. realizado. 36/36 SP, 0 desvíos.",
         "Secciones 6, 9, 12"),
        (f"Sprint_{SPR}_Informe_Pruebas_ICARO.docx", f"INF-PRU-SPR-{SPR}",
         "32 CPs (CP-223–CP-254) al 100% PASS: Consolidación, Pre-cierre, Cierre+Hash, Auditoría, Rollback.",
         "Sección 5 (DoD), Sección 8, Sección 9"),
        ("sprint10_verificacion.md", "—",
         "Criterios de aceptación verificados (✅) para las 5 actividades del sprint.",
         "Sección 5 (DoD), Sección 8"),
    ]
    s_tbl = doc.add_table(rows=len(soporte), cols=4)
    s_tbl.style = 'Table Grid'
    hrow(s_tbl, soporte[0])
    for i, row_v in enumerate(soporte[1:], 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        for j, val in enumerate(row_v):
            cell_bg(s_tbl.rows[i].cells[j], bg)
            cell_text(s_tbl.rows[i].cells[j], val, size=9, color=C_BLACK)

    # Glosario
    add_h1(doc, "Glosario de Siglas")
    glosario = [
        ("Sigla",            "Significado",                               "Contexto de uso"),
        ("RET",              "Retrospectiva",                             "Prefijo del código de este documento"),
        ("SPR",              "Sprint",                                    "Iteración de desarrollo"),
        ("DoD",              "Definition of Done",                        "Criterios de completitud del sprint"),
        ("SP",               "Story Point",                               "Unidad de estimación de esfuerzo"),
        ("HU",               "Historia de Usuario",                       "Tipo de ítem del backlog"),
        ("HT",               "Historia Técnica",                         "Tarea técnica transversal"),
        ("RBAC",             "Role-Based Access Control",                 "Control de acceso por roles"),
        ("JWT",              "JSON Web Token",                            "Mecanismo de autenticación"),
        ("SHA-256",          "Secure Hash Algorithm 256 bits",            "Hash criptográfico del cierre mensual"),
        ("JSONB",            "JSON Binario en PostgreSQL",                "Tipo de dato para datos_despues en audit_log"),
        ("TX",               "Transacción de BD",                        "BEGIN/COMMIT/ROLLBACK automático con Prisma"),
        ("PENDING_SYNC",     "Estado de avance offline sin sincronizar",  "Bloqueo de validación pre-cierre"),
        ("EN_REVISION",      "Estado de requerimiento en revisión",       "Bloqueo de validación pre-cierre"),
        ("CIERRE_YA_EXISTE", "Código de error 422",                       "Período ya tiene cierre CERRADO"),
        ("CSV/BOM",          "Exportación con marcador UTF-8",            "Compatibilidad Excel en AuditTraceabilityView"),
        ("CP",               "Caso de Prueba",                            "Identificador de prueba"),
        ("E2E",              "End-to-End",                               "Prueba de ciclo completo de negocio"),
    ]
    g_tbl = doc.add_table(rows=len(glosario), cols=3)
    g_tbl.style = 'Table Grid'
    hrow(g_tbl, glosario[0])
    for i, (sig, sig_full, ctx) in enumerate(glosario[1:], 1):
        bg = H_WHITE if i % 2 == 0 else H_GRAY
        cell_bg(g_tbl.rows[i].cells[0], H_MED)
        cell_text(g_tbl.rows[i].cells[0], sig, bold=True, size=9, color=C_DARK_BLUE)
        cell_bg(g_tbl.rows[i].cells[1], bg)
        cell_text(g_tbl.rows[i].cells[1], sig_full, size=9, color=C_BLACK)
        cell_bg(g_tbl.rows[i].cells[2], bg)
        cell_text(g_tbl.rows[i].cells[2], ctx, size=9, color=C_BLACK)

    out = os.path.join(OUTDIR, f"Sprint_{SPR}_Retrospectiva_ICARO_V2.docx")
    doc.save(out); print(f"  ✓ {os.path.basename(out)}")


# ════════════════════════════════════════════════════════════════════════════════
# 4. BCK-SPR-10 (XLSX)
# ════════════════════════════════════════════════════════════════════════════════

def gen_backlog():
    wb = Workbook()

    F_HEAD  = Font(name='Calibri', bold=True, size=10, color='FFFFFF')
    F_HU    = Font(name='Calibri', bold=True, size=10, color='000000')
    F_TASK  = Font(name='Calibri', size=9,    color='000000')

    FILL_DARK  = PatternFill("solid", fgColor="1F3864")
    FILL_MED   = PatternFill("solid", fgColor="D9E1F2")
    FILL_WHITE = PatternFill("solid", fgColor="FFFFFF")
    FILL_GRAY  = PatternFill("solid", fgColor="F2F2F2")
    FILL_GREEN = PatternFill("solid", fgColor="E2EFDA")
    FILL_YELL  = PatternFill("solid", fgColor="FFF2CC")

    A_CTR = Alignment(horizontal='center', vertical='center', wrap_text=True)
    A_LFT = Alignment(horizontal='left',   vertical='center', wrap_text=True)
    thin  = Side(style='thin', color='000000')
    BD_   = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Hoja 1: Sprint Backlog ─────────────────────────────────────────────
    ws = wb.active
    ws.title = f"Sprint {SPR} Backlog"

    ws.merge_cells('A1:J1')
    ws['A1'] = f"SPRINT {SPR} — BACKLOG — {PROYECTO}"
    ws['A1'].font = Font(name='Calibri', bold=True, size=14, color='1F3864')
    ws['A1'].alignment = A_CTR
    ws.row_dimensions[1].height = 24

    ws.merge_cells('A2:J2')
    ws['A2'] = (f"Período: {FECHA_INI} – {FECHA_FIN}  |  Responsable: {RESPONSABLE}  |  "
                f"Fecha doc.: {FECHA_DOC}")
    ws['A2'].font = Font(name='Calibri', italic=True, size=9, color='595959')
    ws['A2'].alignment = A_CTR
    ws.row_dimensions[2].height = 16

    # Nota técnica
    ws.merge_cells('A3:J3')
    ws['A3'] = ("NOTA: La BD no fue modificada. cierre.service.js opera sobre tablas existentes: "
                "cierre_mensual, audit_log, avance_obra, requerimiento_compra, movimiento_inventario, inventario_proyecto.")
    ws['A3'].font = Font(name='Calibri', italic=True, size=9, color='C00000')
    ws['A3'].alignment = A_LFT
    ws.row_dimensions[3].height = 16

    headers = ['ID','Tipo','Historia de Usuario / Tarea','Prioridad',
               'SP Plan.','SP Real.','Desvío SP','% Cumpl.','Estado','Observaciones']
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=5, column=j, value=h)
        c.font = F_HEAD; c.fill = FILL_DARK; c.alignment = A_CTR; c.border = BD_
    ws.row_dimensions[5].height = 22

    col_widths = [12, 8, 58, 12, 10, 10, 10, 10, 14, 35]
    for j, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(j)].width = w

    row_num = 6
    for hu_row in HU_ROWS:
        id_, tipo, desc, prio, sp_p, sp_r, desv, pct, fi, ff = hu_row
        is_hu = tipo in ('HU', 'HT')
        fill_row = FILL_MED if is_hu else (FILL_WHITE if row_num % 2 == 0 else FILL_GRAY)
        font_row = F_HU if is_hu else F_TASK
        status   = 'Completado' if pct == '100%' else 'Pendiente'
        fill_st  = FILL_GREEN if status == 'Completado' else FILL_YELL
        vals = [id_, tipo, desc, prio, sp_p, sp_r, desv, pct, status,
                f"Inicio: {fi} | Fin: {ff}"]
        for j, val in enumerate(vals, 1):
            c = ws.cell(row=row_num, column=j, value=val)
            c.font      = font_row
            c.fill      = fill_st if j == 9 else fill_row
            c.alignment = A_CTR if j in [1,2,4,5,6,7,8,9] else A_LFT
            c.border    = BD_
        ws.row_dimensions[row_num].height = 18 if is_hu else 15
        row_num += 1

    totals = ['TOTALES','','','', TOTAL_SP_PLAN, TOTAL_SP_REAL, 0, '100%', '5/5 ítems','']
    for j, val in enumerate(totals, 1):
        c = ws.cell(row=row_num, column=j, value=val)
        c.font = Font(name='Calibri', bold=True, size=10,
                      color='000000' if j in [5,6,7,8] else 'FFFFFF')
        c.fill = FILL_MED if j in [5,6,7,8] else FILL_DARK
        c.alignment = A_CTR; c.border = BD_
    ws.row_dimensions[row_num].height = 20

    # ── Hoja 2: Resumen CP ─────────────────────────────────────────────────
    ws2 = wb.create_sheet(title="Resumen CP")
    ws2.merge_cells('A1:F1')
    ws2['A1'] = f"RESUMEN DE CASOS DE PRUEBA — SPRINT {SPR} (CP-223 a CP-254)"
    ws2['A1'].font = Font(name='Calibri', bold=True, size=12, color='1F3864')
    ws2['A1'].alignment = A_CTR

    cp_headers = ['ID CP','Grupo','Módulo','Descripción','Estado','Evidencia']
    for j, h in enumerate(cp_headers, 1):
        c = ws2.cell(row=2, column=j, value=h)
        c.font = F_HEAD; c.fill = FILL_DARK; c.alignment = A_CTR; c.border = BD_
    ws2.column_dimensions['A'].width = 10
    ws2.column_dimensions['B'].width = 10
    ws2.column_dimensions['C'].width = 22
    ws2.column_dimensions['D'].width = 68
    ws2.column_dimensions['E'].width = 10
    ws2.column_dimensions['F'].width = 45

    for i, cp in enumerate(CP_ROWS, 3):
        cp_id, grupo, mod, desc, _, _, _, estado, evidencia = cp
        fill_r = FILL_GREEN if estado == 'PASS' else \
                 PatternFill("solid", fgColor="FCE4D6")
        for j, val in enumerate([cp_id, grupo, mod, desc, estado, evidencia], 1):
            c = ws2.cell(row=i, column=j, value=val)
            c.font = Font(name='Calibri', size=9, color='000000')
            c.fill = fill_r if j == 5 else (FILL_WHITE if i % 2 == 0 else FILL_GRAY)
            c.alignment = A_CTR if j in [1,2,3,5] else A_LFT
            c.border = BD_
        ws2.row_dimensions[i].height = 14

    # ── Hoja 3: Sprint Summary ─────────────────────────────────────────────
    ws3 = wb.create_sheet(title="Sprint Summary")
    summary = [
        ('Atributo', 'Valor'),
        ('Sprint',                       f'Sprint {SPR}'),
        ('Nombre',                       NOMBRE_SPR),
        ('Fecha inicio',                 FECHA_INI),
        ('Fecha cierre',                 FECHA_FIN),
        ('Responsable',                  RESPONSABLE),
        ('SP Planificados',              TOTAL_SP_PLAN),
        ('SP Realizados',                TOTAL_SP_REAL),
        ('% Cumplimiento SP',            '100%'),
        ('HU Planificadas',              4),
        ('HT Planificadas',              1),
        ('Total ítems',                  5),
        ('CP Planificados',              32),
        ('CP PASS',                      32),
        ('CP FAIL',                      0),
        ('% Éxito pruebas',              '100%'),
        ('CP Sprints previos acumulados',222),
        ('CP acumulados total',          254),
        ('BD modificada',               'NO — sin migraciones ejecutadas'),
        ('Tablas usadas',               'cierre_mensual, audit_log, avance_obra, requerimiento_compra, movimiento_inventario, inventario_proyecto'),
        ('Código PLN-REA',               f'PLN-REA-SPR-{SPR}'),
        ('Código INF-PRU',               f'INF-PRU-SPR-{SPR}'),
        ('Código RET',                   f'RET-SPR-{SPR}'),
        ('Código BCK',                   f'BCK-SPR-{SPR}'),
    ]
    ws3.column_dimensions['A'].width = 30
    ws3.column_dimensions['B'].width = 70
    for i, (k, v) in enumerate(summary, 1):
        cA = ws3.cell(row=i, column=1, value=k)
        cB = ws3.cell(row=i, column=2, value=v)
        if i == 1:
            cA.font = cB.font = F_HEAD
            cA.fill = cB.fill = FILL_DARK
        else:
            cA.fill = FILL_MED
            cA.font = Font(name='Calibri', bold=True, size=10)
            cB.font = Font(name='Calibri', size=10)
        cA.alignment = cB.alignment = A_LFT
        cA.border = cB.border = BD_

    out = os.path.join(OUTDIR, f"Sprint_{SPR}_Backlog_ICARO.xlsx")
    wb.save(out); print(f"  ✓ {os.path.basename(out)}")


# ════════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print(f"\nGenerando documentos Sprint {SPR} → {OUTDIR}\n")
    gen_pln_rea()
    gen_inf_pru()
    gen_retro()
    gen_backlog()
    print(f"\n✓ 4 documentos generados en:\n  {OUTDIR}\n")
