from docx import Document
from openpyxl import load_workbook
import os

BASE = r"c:\Users\Hp\Desktop\Sistema_Gestion_Integral_Constructora\Documentos Base"

def show_doc(path, label):
    doc = Document(path)
    print(f"\n{'='*70}")
    print(f"ARCHIVO: {label}")
    print("="*70)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            runs_info = []
            for r in para.runs:
                info = []
                if r.bold: info.append("bold")
                if r.italic: info.append("italic")
                if r.font.size: info.append(f"sz={int(r.font.size.pt)}")
                if r.font.color and r.font.color.type is not None:
                    try: info.append(f"col={r.font.color.rgb}")
                    except: pass
                if info:
                    runs_info.append(f"[{','.join(info)}]")
            style_info = f"style={para.style.name!r}"
            align_info = f"align={para.alignment}"
            run_str = " ".join(runs_info)
            print(f"  P{i:02d} {style_info} {align_info} | {run_str} | TEXT: {para.text[:100]}")
    print("  --- TABLAS ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"  [T{t_idx}] {len(table.rows)} filas x {len(table.columns)} cols | style={table.style.name!r}")
        for r_idx, row in enumerate(table.rows):
            vals = []
            for c in row.cells:
                # Get bg color
                bg = ""
                try:
                    from docx.oxml.ns import qn
                    shd = c._tc.find(".//" + "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                    if shd is not None:
                        fill = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                        if fill: bg = f"[bg:{fill}]"
                except: pass
                text = c.text[:35].replace("\n", "|")
                # Font info from first run
                finfo = ""
                try:
                    p0 = c.paragraphs[0]
                    if p0.runs:
                        r0 = p0.runs[0]
                        fi = []
                        if r0.bold: fi.append("B")
                        if r0.font.size: fi.append(f"{int(r0.font.size.pt)}pt")
                        try:
                            if r0.font.color.rgb: fi.append(str(r0.font.color.rgb))
                        except: pass
                        if fi: finfo = f"({','.join(fi)})"
                except: pass
                vals.append(f"{bg}{finfo}{text!r}")
            print(f"    R{r_idx:02d}: {vals}")

def show_xlsx(path, label):
    wb = load_workbook(path)
    print(f"\n{'='*70}")
    print(f"XLSX: {label}")
    print("="*70)
    for ws in wb.worksheets:
        print(f"  SHEET: {ws.title}")
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 50)):
            row_data = []
            for cell in row:
                val = str(cell.value)[:30] if cell.value is not None else ""
                bg = ""
                try:
                    if cell.fill and cell.fill.fgColor and cell.fill.fgColor.type == "rgb":
                        bg = f"[{cell.fill.fgColor.rgb}]"
                except: pass
                font_info = ""
                try:
                    fi = []
                    if cell.font.bold: fi.append("B")
                    if cell.font.size: fi.append(f"{cell.font.size}pt")
                    if cell.font.color and cell.font.color.type == "rgb": fi.append(cell.font.color.rgb)
                    if fi: font_info = f"({','.join(str(x) for x in fi)})"
                except: pass
                row_data.append(f"[{cell.coordinate}]{bg}{font_info}{val!r}")
            non_empty = [x for x in row_data if "''" not in x or "[" in x]
            if non_empty:
                print(f"    {'  '.join(non_empty[:6])}")

for sprint in ["1", "2", "3"]:
    folder = f"Sprint {sprint} documentacion"
    sp_pad = sprint.zfill(2)
    v_suffix = " V1" if sprint == "1" else ""
    v2_suffix = " V1" if sprint == "1" else ""
    
    # PLN-REA
    pln = os.path.join(BASE, folder, f"Sprint_{sp_pad}_Planificado_Realizado_ICARO{v_suffix}.docx")
    if os.path.exists(pln):
        show_doc(pln, f"SPR{sp_pad} PLN-REA")
    
    # INF-PRU
    inf = os.path.join(BASE, folder, f"Sprint_{sp_pad}_Informe_Pruebas_ICARO{v_suffix}.docx")
    if os.path.exists(inf):
        show_doc(inf, f"SPR{sp_pad} INF-PRU")
    
    # RETRO V2
    ret = os.path.join(BASE, folder, f"Sprint_{sp_pad}_Retrospectiva_ICARO_V2.docx")
    if os.path.exists(ret):
        show_doc(ret, f"SPR{sp_pad} RETRO V2")
    
    # BACKLOG
    bck = os.path.join(BASE, folder, f"Sprint_{sp_pad}_Backlog_ICARO.xlsx")
    if os.path.exists(bck):
        show_xlsx(bck, f"SPR{sp_pad} BACKLOG")
