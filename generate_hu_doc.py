"""
Script principal para generar el documento Historias_Actualizadas.docx
con todas las 26 historias de usuario en el mismo formato del template.
"""
import os
import re
import html
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ============================================================
# CONFIG
# ============================================================
DRAWIO_FOLDER = r'c:\Users\Hp\Desktop\Sistema_Gestion_Integral_Constructora\Documentos Version 2.0\HISTORIAS DE USUARIO V2'
TEMPLATE_PATH = r'c:\Users\Hp\Desktop\Sistema_Gestion_Integral_Constructora\Documentos Version 2.0\Historias_Usuario_Tecnica V2.docx'
OUTPUT_PATH   = r'c:\Users\Hp\Desktop\Sistema_Gestion_Integral_Constructora\Documentos Version 2.0\Historias_Actualizadas.docx'

# Colors from template
BORDER_COLOR  = '3A2A0F'
HEADER_FILL   = 'F3EFE6'
FONT_NAME     = 'Arial'
FONT_SZ_LABEL = Pt(10.5)   # 133350 EMU ≈ 10.5 pt
FONT_SZ_BODY  = Pt(10)     # 127000 EMU ≈ 10 pt
LEFT_COL_W    = 2325       # dxa (twips)

# ============================================================
# PARSE DRAWIO FILES
# ============================================================

def clean_html(raw: str) -> str:
    """Remove HTML tags and decode entities."""
    if not raw:
        return ''
    # Add newlines before block-closing tags so split works
    raw = re.sub(r'</?(p|div|br|li)[^>]*>', '\n', raw, flags=re.IGNORECASE)
    # Remove remaining HTML tags
    text = re.sub(r'<[^>]+>', '', raw)
    # Decode HTML entities
    text = html.unescape(text)
    # Normalize spaces within lines, keep newlines
    lines = [re.sub(r'[ \t]+', ' ', l).strip() for l in text.split('\n')]
    text = '\n'.join(l for l in lines if l)
    return text.strip()


def extract_bullet_points(raw: str) -> list:
    """Extract bullet point lines from HTML content (handles <p> and <div>)."""
    if not raw:
        return []
    # Normalize block tags to newlines
    normalized = re.sub(r'</?(p|div|br|li)[^>]*>', '\n', raw, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r'<[^>]+>', '', normalized)
    text = html.unescape(text)
    bullets = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Remove leading bullet/dash/arrow chars
        line = re.sub(r'^[•·\-–→]\s*', '', line).strip()
        if line and len(line) > 5:
            bullets.append(line)
    return bullets


def parse_drawio(filepath: str) -> dict:
    """Parse a .drawio XML file and extract user story data."""
    tree = etree.parse(filepath)
    root = tree.getroot()

    data = {
        'id': '',
        'title': '',
        'description': '',
        'estimation': '',
        'priority': '',
        'dependencies': '',
        'criteria': []
    }

    # Get all mxCell elements with value attribute
    cells = root.findall('.//mxCell[@value]')

    for cell in cells:
        value = cell.get('value', '').strip()
        if not value:
            continue

        style = cell.get('style', '')
        v_clean = clean_html(value)

        # Detect HU ID (e.g., HU-01)
        if re.match(r'^HU-\d+$', v_clean.strip()):
            data['id'] = v_clean.strip()
            continue

        # Also check for HTML-wrapped HU-XX
        if re.search(r'HU-\d+', v_clean) and len(v_clean) < 20 and 'Como' not in v_clean:
            match = re.search(r'HU-\d+', v_clean)
            if match and not data['id']:
                data['id'] = match.group(0)
                continue

        # Detect title (bold, short, no "Como")
        if '<b>' in value and 'Como' not in v_clean and len(v_clean) < 160:
            title_match = re.search(r'<b[^>]*>(.*?)</b>', value, re.DOTALL)
            if title_match:
                candidate = clean_html(title_match.group(1)).strip()
                # Titles should be single-line — collapse newlines
                candidate = re.sub(r'\s+', ' ', candidate).strip()
                if candidate and 'HU-' not in candidate and len(candidate) > 5:
                    data['title'] = candidate
                    continue

        # Detect description (Como... quiero... para...)
        if re.search(r'\bComo\b', v_clean) and re.search(r'\bquiero\b', v_clean):
            data['description'] = v_clean
            continue

        # Detect estimation
        if re.search(r'[Ee]stima[a-zA-Záéíóú]*\s*(tiempo|de esfuerzo)?.*:\s*\d+', v_clean):
            data['estimation'] = v_clean
            continue

        # Detect priority
        if re.search(r'[Pp]rioridad\s*:', v_clean) and not re.search(r'[Ee]stima', v_clean):
            data['priority'] = re.sub(r'[Pp]rioridad\s*:?\s*', '', v_clean).strip()
            continue

        # Detect dependencies (require colon to avoid matching words like "independientes")
        if re.search(r'[Dd]ependien[a-z]*\s*:', v_clean):
            dep_text = re.sub(r'[Dd]ependien[a-z]*\s*:?\s*', '', v_clean).strip()
            # Clean extra spaces and normalize
            dep_text = re.sub(r'\s+', ' ', dep_text).strip()
            data['dependencies'] = dep_text if dep_text else 'Sin dependencias'
            continue

        # Detect acceptance criteria (bullet points)
        # Usually in a separate cell with multiple <p> bullet points
        if value.count('<p') >= 3 or (v_clean.count('•') >= 2) or ('Con ' in v_clean and 'cuando' in v_clean):
            if len(v_clean) > 200:
                data['criteria'] = extract_bullet_points(value)
                continue

    # If criteria still empty, try to find a cell with many bullet chars
    if not data['criteria']:
        for cell in cells:
            value = cell.get('value', '').strip()
            v_clean = clean_html(value)
            if len(v_clean) > 200 and ('cuando' in v_clean or 'el sistema' in v_clean.lower()):
                data['criteria'] = extract_bullet_points(value)
                break

    # Fallback: extract estimation from combined field if needed
    if not data['estimation'] and not data['priority']:
        for cell in cells:
            v = clean_html(cell.get('value', ''))
            if 'horas' in v.lower():
                data['estimation'] = v
            if 'prioridad' in v.lower() and ':' in v:
                parts = re.split(r'[Pp]rioridad\s*:?\s*', v)
                if len(parts) > 1:
                    data['priority'] = parts[1].strip()

    return data


# ============================================================
# LOAD ALL DRAWIO FILES
# ============================================================

drawio_files = sorted([
    f for f in os.listdir(DRAWIO_FOLDER) if f.endswith('.drawio')
])

stories = []
for filename in drawio_files:
    filepath = os.path.join(DRAWIO_FOLDER, filename)
    story = parse_drawio(filepath)
    # Fallback ID from filename
    if not story['id']:
        m = re.match(r'(HU-\d+)', filename)
        if m:
            story['id'] = m.group(1)
    # Fallback title from filename
    if not story['title']:
        m = re.match(r'HU-\d+ [·.] (.+)\.drawio', filename)
        if m:
            story['title'] = m.group(1)
    stories.append(story)
    # Normalize title (collapse any residual newlines)
    stories[-1]['title'] = re.sub(r'\s+', ' ', stories[-1]['title']).strip()
    print(f"  Parsed {story['id']}: '{story['title'][:50]}' | criteria={len(story['criteria'])} | deps={story['dependencies'][:40]}")

print(f"\nTotal stories parsed: {len(stories)}")

# ============================================================
# DOCX HELPERS
# ============================================================

def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '12')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), BORDER_COLOR)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_cell_fill(cell, fill_color: str):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)


def set_cell_width(cell, width_dxa: int):
    """Set cell width in dxa units."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def set_table_borders(table):
    """Set borders on a table (not just cells)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '12')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), BORDER_COLOR)
        tblBorders.append(tag)
    # Remove old borders if exists
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)

    # Fixed layout
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    old_layout = tblPr.find(qn('w:tblLayout'))
    if old_layout is not None:
        tblPr.remove(old_layout)
    tblPr.append(tblLayout)


def add_run(para, text: str, bold=False, size=None, color=None, font_name=None):
    """Add a run to a paragraph with formatting."""
    run = para.add_run(text)
    run.font.name = font_name or FONT_NAME
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_label_cell(cell, label: str):
    """Format the left (label) cell of the table."""
    set_cell_fill(cell, HEADER_FILL)
    set_cell_width(cell, LEFT_COL_W)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(para, label, bold=True, size=FONT_SZ_LABEL)


def add_value_cell_text(cell, text: str, font_size=None):
    """Add plain text to the right (value) cell."""
    para = cell.paragraphs[0]
    para.clear()
    add_run(para, text, size=font_size or FONT_SZ_BODY)


def add_value_cell_bullets(cell, bullets: list):
    """Add bullet points to value cell."""
    if not bullets:
        cell.paragraphs[0].clear()
        add_run(cell.paragraphs[0], '(Sin criterios definidos)', size=FONT_SZ_BODY)
        return
    first = True
    for bullet in bullets:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        add_run(para, f'• {bullet}', size=FONT_SZ_BODY)
        para.paragraph_format.space_after = Pt(2)


def format_all_cells(table):
    """Apply border formatting to all cells in table."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def add_hu_table(doc, story: dict):
    """Add a complete HU table to the document."""
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Normal'
    set_table_borders(table)

    # Row 0: Título
    add_label_cell(table.rows[0].cells[0], '1. Título')
    add_value_cell_text(table.rows[0].cells[1],
                        f"{story['id']} – {story['title']}", FONT_SZ_LABEL)
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    # Row 1: Descripción
    add_label_cell(table.rows[1].cells[0], '\n2. Descripción')
    add_value_cell_text(table.rows[1].cells[1], story['description'], FONT_SZ_BODY)

    # Row 2: Criterios de Aceptación
    add_label_cell(table.rows[2].cells[0], '\n\n3. Criterios de Aceptación')
    add_value_cell_bullets(table.rows[2].cells[1], story['criteria'])

    # Row 3: Tareas
    add_label_cell(table.rows[3].cells[0], '\n4. Tareas (si es necesario)')
    add_value_cell_text(table.rows[3].cells[1],
        'Revisar e integrar con módulos dependientes según criterios de aceptación.',
        FONT_SZ_BODY)

    # Row 4: Pruebas
    add_label_cell(table.rows[4].cells[0], '\n5. Pruebas (si es necesario)')
    add_value_cell_text(table.rows[4].cells[1],
        'Prueba funcional de cada criterio de aceptación descrito arriba.',
        FONT_SZ_BODY)

    # Row 5: Dependencias
    add_label_cell(table.rows[5].cells[0], '6. Dependencias')
    deps = story['dependencies'] if story['dependencies'] else 'Sin dependencias'
    add_value_cell_text(table.rows[5].cells[1], deps, FONT_SZ_LABEL)

    # Row 6: Estimación / Prioridad
    add_label_cell(table.rows[6].cells[0], '7. Estimación de\nEsfuerzo / 8. Prioridad')
    est = story['estimation'] if story['estimation'] else 'No especificada'
    pri = story['priority'] if story['priority'] else 'Alta'
    add_value_cell_text(table.rows[6].cells[1],
                        f"{est}    Prioridad: {pri}", FONT_SZ_LABEL)

    # Apply border to all cells
    format_all_cells(table)


# ============================================================
# BUILD THE DOCUMENT
# ============================================================

# Start from a fresh copy of the template
doc = Document(TEMPLATE_PATH)

# Find and add/remove existing HU section if present, then append
# Since template only has HT, we just append a new HU section at the end

# Add page break
doc.add_page_break()

# Add main heading
heading = doc.add_paragraph('Historias de Usuario (Versión 2.0)', style='Heading 1')

# Process each story
for i, story in enumerate(stories):
    # Add HU sub-heading
    sub = doc.add_paragraph(style='Heading 2' if 'Heading 2' in [s.name for s in doc.styles] else 'Normal')
    sub.clear()
    run = sub.add_run(f"{story['id']}. {story['title']}")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    sub.paragraph_format.space_before = Pt(12)
    sub.paragraph_format.space_after = Pt(4)

    # Add the table
    add_hu_table(doc, story)

    # Space after table
    spacer = doc.add_paragraph('')
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)

    print(f"  Added table for {story['id']}")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"\n✅ Documento generado: {OUTPUT_PATH}")
